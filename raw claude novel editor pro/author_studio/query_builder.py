"""
query_builder.py
================
Generates professional query submission packages as formatted .docx files.

Produces four documents that authors currently pay $200–800 for from
professional query consultants:

  1. Query Letter     — Standard one-page agent query letter
  2. One-Page Synopsis — The complete story in ~500 words
  3. Three-Page Synopsis — Detailed narrative summary (~1,200 words)
  4. Author Bio Sheet  — Standalone bio page for submissions

All templates follow current AAR (Association of Authors' Representatives)
and Publisher's Marketplace guidelines.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, Dict


# ── Query Package Data Model ──────────────────────────────────────────────────

class QueryPackageData:
    """
    All the information needed to generate the complete submission package.
    The app collects these from the author via form fields.
    """
    def __init__(
        self,
        # Manuscript identity
        title:           str,
        author_name:     str,
        genre:           str,
        word_count:      int,
        series_note:     str  = "",   # e.g. "Standalone with series potential"

        # Author contact
        email:           str  = "",
        phone:           str  = "",
        address:         str  = "",
        website:         str  = "",

        # Publishing credentials
        bio_credits:     str  = "",   # Previous publications, awards, relevant experience
        bio_platform:    str  = "",   # Social media, newsletter, professional platform

        # Comparable titles (comps)
        comp_1_title:    str  = "",
        comp_1_author:   str  = "",
        comp_1_year:     str  = "",
        comp_2_title:    str  = "",
        comp_2_author:   str  = "",
        comp_2_year:     str  = "",

        # Story elements
        protagonist:     str  = "",   # Name and defining trait
        setting:         str  = "",   # Time and place
        inciting_event:  str  = "",   # What kicks off the story
        central_conflict: str = "",   # The main dramatic question
        stakes:          str  = "",   # What happens if the protagonist fails
        theme:           str  = "",   # Thematic core (one sentence)

        # Synopsis content
        synopsis_plot:   str  = "",   # Full plot summary the author writes (we format it)
    ):
        self.title            = title
        self.author_name      = author_name
        self.genre            = genre
        self.word_count       = word_count
        self.series_note      = series_note
        self.email            = email
        self.phone            = phone
        self.address          = address
        self.website          = website
        self.bio_credits      = bio_credits
        self.bio_platform     = bio_platform
        self.comp_1_title     = comp_1_title
        self.comp_1_author    = comp_1_author
        self.comp_1_year      = comp_1_year
        self.comp_2_title     = comp_2_title
        self.comp_2_author    = comp_2_author
        self.comp_2_year      = comp_2_year
        self.protagonist      = protagonist
        self.setting          = setting
        self.inciting_event   = inciting_event
        self.central_conflict = central_conflict
        self.stakes           = stakes
        self.theme            = theme
        self.synopsis_plot    = synopsis_plot

    def has_comps(self) -> bool:
        return bool(self.comp_1_title and self.comp_1_author)

    def comp_string(self) -> str:
        parts = []
        if self.comp_1_title and self.comp_1_author:
            yr = f" ({self.comp_1_year})" if self.comp_1_year else ""
            parts.append(f"{self.comp_1_title}{yr} by {self.comp_1_author}")
        if self.comp_2_title and self.comp_2_author:
            yr = f" ({self.comp_2_year})" if self.comp_2_year else ""
            parts.append(f"{self.comp_2_title}{yr} by {self.comp_2_author}")
        if len(parts) == 2:
            return f"{parts[0]} meets {parts[1]}"
        return parts[0] if parts else ""


# ── Document Builder Base ─────────────────────────────────────────────────────

class _DocBuilder:
    """Shared formatting utilities for all query documents."""

    FONT    = "Times New Roman"
    SIZE    = Pt(12)
    MARGIN  = Inches(1.0)

    def _new_doc(self) -> Document:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = self.MARGIN
            section.bottom_margin = self.MARGIN
            section.left_margin   = self.MARGIN
            section.right_margin  = self.MARGIN
        # Set Normal style
        normal = doc.styles["Normal"]
        normal.font.name = self.FONT
        normal.font.size = self.SIZE
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(0)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        return doc

    def _p(
        self, doc: Document, text: str = "",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        bold: bool = False,
        italic: bool = False,
        size: Optional[Pt] = None,
        single_space: bool = False,
        space_after: Pt = Pt(0),
    ):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = space_after
        if single_space:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            run = para.add_run(text)
            run.font.name   = self.FONT
            run.font.size   = size or self.SIZE
            run.bold        = bold
            run.italic      = italic
        return para

    def _indent_p(self, doc: Document, text: str):
        """Body paragraph with 0.5-inch first-line indent."""
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if text:
            run = para.add_run(text)
            run.font.name = self.FONT
            run.font.size = self.SIZE
        return para


# ── Query Letter Builder ──────────────────────────────────────────────────────

class QueryLetterBuilder(_DocBuilder):
    """
    Generates a standard one-page query letter following AAR guidelines.

    Structure (in order):
      1. Author contact block
      2. Date
      3. Agent personalisation placeholder
      4. Hook paragraph (1–2 sentences)
      5. Plot paragraph (3–5 sentences: protagonist, conflict, stakes)
      6. Comp titles paragraph
      7. Manuscript details paragraph
      8. Bio paragraph
      9. Closing
    """

    def build(self, data: QueryPackageData, output_path: str) -> str:
        doc = self._new_doc()

        # Set to single-space for the query letter (industry standard)
        for style in doc.styles:
            if style.name == "Normal":
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        def sp(text="", bold=False, italic=False, after=Pt(0)):
            return self._p(doc, text, bold=bold, italic=italic,
                           single_space=True, space_after=after)

        # ── Contact block ──
        sp(data.author_name, bold=True)
        if data.address:   sp(data.address)
        if data.phone:     sp(data.phone)
        sp(data.email or "your.email@example.com")
        if data.website:   sp(data.website)
        sp()  # blank line

        # ── Agent personalisation ──
        sp("[Agent Name]")
        sp("[Agency Name]")
        sp("[Agency Address]")
        sp()
        sp("Dear [Agent Name],")
        sp()

        # ── Hook paragraph ──
        hook = self._compose_hook(data)
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(12)
        run = para.add_run(hook)
        run.font.name = self.FONT
        run.font.size = self.SIZE

        # ── Plot paragraph ──
        plot = self._compose_plot(data)
        para2 = doc.add_paragraph()
        para2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after  = Pt(12)
        run2 = para2.add_run(plot)
        run2.font.name = self.FONT
        run2.font.size = self.SIZE

        # ── Comps paragraph ──
        if data.has_comps():
            comp_text = (
                f"{data.title} will appeal to readers of {data.comp_string()}. "
                f"Like {data.comp_1_author.split()[-1]}'s work, this novel "
                f"[briefly note the thematic or stylistic parallel — edit this line]."
            )
            para3 = doc.add_paragraph()
            para3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para3.paragraph_format.space_before = Pt(0)
            para3.paragraph_format.space_after  = Pt(12)
            para3.add_run(comp_text).font.name = self.FONT

        # ── Manuscript details ──
        series = f" {data.series_note}." if data.series_note else ""
        ms_text = (
            f"{data.title.upper()} is a {data.genre} novel complete at "
            f"{data.word_count:,} words.{series} "
            f"The full manuscript is available upon request."
        )
        para4 = doc.add_paragraph()
        para4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para4.paragraph_format.space_before = Pt(0)
        para4.paragraph_format.space_after  = Pt(12)
        para4.add_run(ms_text).font.name = self.FONT

        # ── Bio ──
        bio_text = self._compose_bio(data)
        para5 = doc.add_paragraph()
        para5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para5.paragraph_format.space_before = Pt(0)
        para5.paragraph_format.space_after  = Pt(12)
        para5.add_run(bio_text).font.name = self.FONT

        # ── Closing ──
        sp("Thank you for your time and consideration.")
        sp()
        sp("Sincerely,")
        sp()
        sp(data.author_name)

        # ── Footer note ──
        sp()
        sp(
            "[NOTE: Personalise the agent salutation, the comp paragraph, "
            "and the bio paragraph before sending. Delete this line.]",
            italic=True,
        )

        doc.save(output_path)
        return output_path

    def _compose_hook(self, data: QueryPackageData) -> str:
        if data.protagonist and data.central_conflict:
            return (
                f"When {data.protagonist} {data.inciting_event or '[inciting event]'}, "
                f"they must {data.central_conflict or '[central conflict]'} — "
                f"or {data.stakes or '[stakes: what is lost if they fail]'}."
            )
        return (
            f"[Write your hook sentence here. One to two sentences that capture "
            f"the protagonist, the inciting event, and the central dramatic question. "
            f"This is the most important sentence in the query — it must compel the agent to read on.]"
        )

    def _compose_plot(self, data: QueryPackageData) -> str:
        if data.synopsis_plot and len(data.synopsis_plot) > 100:
            # Use the first 400 chars of the author's synopsis as plot para foundation
            trimmed = data.synopsis_plot[:500].rsplit(" ", 1)[0]
            return trimmed + " [Continue: what is the central conflict and what are the stakes?]"
        return (
            f"[Write 3–5 sentences here describing: (1) your protagonist and their "
            f"situation at the story's opening; (2) the inciting event that disrupts "
            f"their world; (3) the central conflict and the obstacles they face; "
            f"(4) the emotional and external stakes. Do not reveal the ending in a query letter.]"
        )

    def _compose_bio(self, data: QueryPackageData) -> str:
        name = data.author_name
        parts = []
        if data.bio_credits:
            parts.append(data.bio_credits)
        if data.bio_platform:
            parts.append(data.bio_platform)

        if parts:
            return f"{name} is {'. '.join(parts)}."
        return (
            f"{name} [add your publishing credits, relevant professional experience, "
            f"education, or platform here. If this is your debut novel with no prior "
            f"credits, simply state: 'This is my debut novel.' Agents do not penalise debut authors.]"
        )


# ── Synopsis Builder ──────────────────────────────────────────────────────────

class SynopsisBuilder(_DocBuilder):
    """
    Generates formatted one-page and three-page synopses.
    The synopsis is the author's plot summary — this class formats it correctly
    and adds the professional structural elements agents expect.
    """

    def build(self, data: QueryPackageData, output_path: str, pages: int = 1) -> str:
        """
        pages: 1 = one-page synopsis (~500 words); 3 = three-page (~1,200 words)
        """
        doc  = self._new_doc()
        t    = "One-Page Synopsis" if pages == 1 else "Three-Page Synopsis"

        # ── Header block ──
        self._p(doc, data.author_name, single_space=True)
        self._p(doc, data.email or "your.email@example.com", single_space=True)
        self._p(doc, "", single_space=True)

        # ── Title block ──
        self._p(doc, data.title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True, single_space=True)
        self._p(doc, t, align=WD_ALIGN_PARAGRAPH.CENTER, single_space=True)
        self._p(doc, f"{data.genre} | {data.word_count:,} words",
                align=WD_ALIGN_PARAGRAPH.CENTER, single_space=True)
        self._p(doc, "", single_space=True)

        # ── Important convention note ──
        note = doc.add_paragraph()
        note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        note.paragraph_format.space_after = Pt(12)
        note_run = note.add_run(
            "[SYNOPSIS CONVENTION: Write in present tense, third person, "
            "even if your novel is in first person. The synopsis must reveal "
            "the ending — agents need to see that the story resolves. "
            "Delete this instruction before sending.]"
        )
        note_run.italic   = True
        note_run.font.name = self.FONT
        note_run.font.size = Pt(10)

        # ── Synopsis body ──
        if data.synopsis_plot:
            # Format the author's synopsis text properly
            paragraphs = [p.strip() for p in data.synopsis_plot.split("\n") if p.strip()]
            first = True
            for para_text in paragraphs:
                if first:
                    # First paragraph: no indent
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    run = p.add_run(para_text)
                    run.font.name = self.FONT
                    run.font.size = self.SIZE
                    first = False
                else:
                    self._indent_p(doc, para_text)
        else:
            # Scaffold
            length_note = (
                "approximately 500 words covering: opening situation → inciting event → "
                "rising action → midpoint → darkest moment → climax → resolution"
                if pages == 1 else
                "approximately 1,200 words covering all major plot beats, character arcs, "
                "and subplot threads, with the ending revealed in full"
            )
            self._p(doc,
                    f"[Write your synopsis here ({length_note}). "
                    f"Paste your full plot summary — this template will format it correctly. "
                    f"Character names should appear in CAPS on first mention only.]",
                    italic=True, single_space=False)

        doc.save(output_path)
        return output_path


# ── Back Matter Builder ───────────────────────────────────────────────────────

class BackMatterBuilder(_DocBuilder):
    """
    Generates professional back matter pages for the formatted manuscript:
    - Copyright page (for self-publishing)
    - About the Author page
    - Also by the Author page (if applicable)
    - Acknowledgements page template
    """

    def build_copyright_page(
        self,
        data: QueryPackageData,
        publisher_name:  str = "",
        publication_year: str = "2025",
        isbn_print: str = "",
        isbn_ebook: str = "",
        output_path: str = "",
    ) -> str:
        import datetime
        year = publication_year or str(datetime.datetime.now().year)
        doc  = self._new_doc()

        def sp(text="", italic=False, size=None, after=Pt(6)):
            self._p(doc, text, italic=italic, size=size,
                    single_space=True, space_after=after)

        # Push to lower third of page
        for _ in range(20):
            self._p(doc, "", single_space=True)

        sp(f"Copyright © {year} {data.author_name}")
        sp("All rights reserved.")
        sp()
        sp(
            "No part of this publication may be reproduced, stored in a retrieval "
            "system, or transmitted in any form or by any means — electronic, "
            "mechanical, photocopy, recording, or any other — without prior written "
            "permission from the publisher, except for brief quotations in reviews.",
            italic=False,
        )
        sp()
        sp("This is a work of fiction. Names, characters, places, and incidents are "
           "either the product of the author's imagination or are used fictitiously. "
           "Any resemblance to actual persons, living or dead, events, or locales "
           "is entirely coincidental.")
        sp()

        if publisher_name:
            sp(f"Published by {publisher_name}")
        if isbn_print:
            sp(f"ISBN (Print): {isbn_print}")
        if isbn_ebook:
            sp(f"ISBN (eBook): {isbn_ebook}")

        sp()
        sp(f"First published {year}")
        sp()
        sp("Printed in [Country]")

        doc.save(output_path)
        return output_path

    def build_about_author(
        self,
        data:         QueryPackageData,
        bio_long:     str = "",
        output_path:  str = "",
    ) -> str:
        doc = self._new_doc()
        self._p(doc, "About the Author", align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True, single_space=True, space_after=Pt(24))

        bio = bio_long or (
            f"{data.author_name} [write your author biography here — 100 to 200 words. "
            f"Include where you live, any relevant professional background, previous publications, "
            f"and a line about your personal life or interests that humanises you to the reader. "
            f"Write in third person.]"
        )

        # Split into paragraphs
        first = True
        for para_text in [p.strip() for p in bio.split("\n") if p.strip()]:
            if first:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                p.paragraph_format.space_after = Pt(0)
                p.add_run(para_text).font.name = self.FONT
                first = False
            else:
                self._indent_p(doc, para_text)

        doc.save(output_path)
        return output_path


# ── Convenience: Build Complete Package ──────────────────────────────────────

def build_full_package(
    data:           QueryPackageData,
    output_dir:     str,
    include_query:  bool = True,
    include_synopsis_1: bool = True,
    include_synopsis_3: bool = True,
    include_back_matter: bool = True,
) -> Dict[str, str]:
    """
    Generates the complete submission package.
    Returns a dict mapping document type → file path.
    """
    import os
    os.makedirs(output_dir, exist_ok=True)

    results = {}
    safe = data.title.replace(" ", "_").replace("/", "-")[:30]

    if include_query:
        path = os.path.join(output_dir, f"{safe}_query_letter.docx")
        QueryLetterBuilder().build(data, path)
        results["query_letter"] = path

    if include_synopsis_1:
        path = os.path.join(output_dir, f"{safe}_synopsis_1page.docx")
        SynopsisBuilder().build(data, path, pages=1)
        results["synopsis_1page"] = path

    if include_synopsis_3:
        path = os.path.join(output_dir, f"{safe}_synopsis_3page.docx")
        SynopsisBuilder().build(data, path, pages=3)
        results["synopsis_3page"] = path

    if include_back_matter:
        path = os.path.join(output_dir, f"{safe}_about_author.docx")
        BackMatterBuilder().build_about_author(data, output_path=path)
        results["about_author"] = path

        path = os.path.join(output_dir, f"{safe}_copyright_page.docx")
        BackMatterBuilder().build_copyright_page(data, output_path=path)
        results["copyright_page"] = path

    return results
