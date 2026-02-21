"""
formatter.py
============
Builds a perfectly formatted manuscript .docx from the parsed paragraph list.
Applies every rule in the selected NovelTemplate with zero ambiguity.

Design principles
-----------------
1. Creates a BRAND NEW document — never modifies the original.
2. Applies typographic rules that mirror professional publishing:
   - First paragraph after a chapter heading has NO first-line indent.
   - All subsequent body paragraphs use the template indent.
   - Scene breaks are centred and spaced symmetrically.
   - Page numbers are real field codes, not static text.
3. The title page is auto-generated and follows submission standards.
4. Every setting comes from the template — nothing is hardcoded here.
"""

import re
from copy import deepcopy
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from templates import NovelTemplate
from parser import ParsedParagraph, PARA_BODY, PARA_CHAPTER, PARA_SCENE_BREAK, PARA_EMPTY, PARA_FRONT_MATTER


# ── Page size constants ───────────────────────────────────────────────────────
_LETTER_W = Inches(8.5)
_LETTER_H = Inches(11.0)
_A4_W     = Cm(21.0)
_A4_H     = Cm(29.7)


class NovelFormatter:
    """
    Takes a list of ParsedParagraph objects and a NovelTemplate,
    and produces a formatted .docx file.

    Usage:
        formatter = NovelFormatter(template, author="Jane Smith", title="The Lost Hours")
        output_path, warnings = formatter.build(parsed_paragraphs, "output.docx")
    """

    def __init__(
        self,
        template:    NovelTemplate,
        author:      str,
        title:       str,
        word_count:  Optional[int] = None,
    ):
        self.template   = template
        self.author     = author.strip()
        self.title      = title.strip()
        self.word_count = word_count

    # ── Public Entry Point ────────────────────────────────────────────────────

    def build(
        self,
        paragraphs:  List[ParsedParagraph],
        output_path: str,
    ) -> Tuple[str, List[str]]:
        """
        Builds and saves the formatted document.
        Returns (output_path, list_of_warnings).
        """
        warnings: List[str] = []
        t = self.template

        doc = Document()
        self._configure_page(doc)
        self._set_default_style(doc)

        if t.include_running_header:
            self._add_running_header(doc)

        self._add_title_page(doc)
        self._add_body(doc, paragraphs, warnings)

        doc.save(output_path)
        return output_path, warnings

    # ── Page Configuration ────────────────────────────────────────────────────

    def _configure_page(self, doc: Document):
        t = self.template
        for section in doc.sections:
            if t.page_size == "a4":
                section.page_width  = _A4_W
                section.page_height = _A4_H
            else:
                section.page_width  = _LETTER_W
                section.page_height = _LETTER_H

            section.top_margin    = Inches(t.margin_top_in)
            section.bottom_margin = Inches(t.margin_bottom_in)
            section.left_margin   = Inches(t.margin_left_in)
            section.right_margin  = Inches(t.margin_right_in)

    # ── Default Document Style ────────────────────────────────────────────────

    def _set_default_style(self, doc: Document):
        """
        Applies font, size, and spacing to the Normal style so that every
        paragraph inherits these settings unless explicitly overridden.
        """
        t = self.template
        normal = doc.styles["Normal"]

        normal.font.name = t.font_name
        normal.font.size = Pt(t.font_size_pt)

        pf = normal.paragraph_format
        pf.space_before = Pt(t.space_before_para_pt)
        pf.space_after  = Pt(t.space_after_para_pt)
        self._set_line_spacing(pf, t.line_spacing)

    def _set_line_spacing(self, pf, spacing: str):
        if spacing == "double":
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        elif spacing == "one_half":
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ── Running Header ────────────────────────────────────────────────────────

    def _add_running_header(self, doc: Document):
        """
        Adds a running header to the default section.
        - First page header is suppressed (title page has no header).
        - Header contains author name, short title, and an auto page-number field.
        """
        t = self.template
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        # Clear any content Word auto-creates
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p._p.getparent().remove(p._p)

        para = header.add_paragraph()
        para.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT if t.header_align_right else WD_ALIGN_PARAGRAPH.LEFT
        )
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        header_size = Pt(max(t.font_size_pt - 2, 9))

        # Resolve {author} and {title} tokens; {page} becomes a field code
        fmt = t.header_format
        parts = re.split(r"(\{page\})", fmt, flags=re.IGNORECASE)
        for part in parts:
            if part.lower() == "{page}":
                self._insert_page_number_field(para, t.font_name, header_size)
            else:
                text = (
                    part
                    .replace("{author}", self.author)
                    .replace("{title}",  self._short_title())
                )
                if text:
                    run = para.add_run(text)
                    run.font.name = t.font_name
                    run.font.size = header_size

    def _short_title(self, max_len: int = 28) -> str:
        t = self.title.upper()
        return t[:max_len] + "\u2026" if len(t) > max_len else t

    def _insert_page_number_field(self, para, font_name: str, font_size: Pt):
        """Inserts a real { PAGE } field code — renders as live page number in Word."""
        run = para.add_run()
        run.font.name = font_name
        run.font.size = font_size

        fld_begin    = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"),    "begin")
        instr        = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
        fld_separate = OxmlElement("w:fldChar"); fld_separate.set(qn("w:fldCharType"), "separate")
        fld_end      = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"),      "end")

        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_separate)
        r.append(fld_end)

    # ── Title Page ────────────────────────────────────────────────────────────

    def _add_title_page(self, doc: Document):
        """
        Generates a properly formatted submission title page.
        Layout follows Association of Authors' Representatives guidelines:
          - Top-left:    Author name and contact details
          - Top-right:   Word count (approximate)
          - Mid-page:    Title (centred, uppercase)
          - Below title: Byline and genre
        """
        t = self.template

        def _p(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size_delta=0):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if text:
                run = p.add_run(text)
                run.font.name = t.font_name
                run.font.size = Pt(t.font_size_pt + size_delta)
                run.bold = bold
            return p

        # Contact block (top-left)
        _p(self.author)
        _p("your.email@example.com")
        _p("Your Address, City, Country")
        _p("Your Phone Number")
        _p()

        # Word count (top-right)
        wc = f"Approx. {self.word_count:,} words" if self.word_count else "Word Count: [add here]"
        _p(wc, align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Blank lines → push title to one-third down the page
        for _ in range(10):
            _p()

        # Title
        _p(self.title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_delta=2)
        _p()
        _p(f"by {self.author}", align=WD_ALIGN_PARAGRAPH.CENTER)
        _p()
        _p("A Novel", align=WD_ALIGN_PARAGRAPH.CENTER)

        # Mandatory page break
        doc.add_page_break()

    # ── Body Content ──────────────────────────────────────────────────────────

    def _add_body(
        self,
        doc:        Document,
        paragraphs: List[ParsedParagraph],
        warnings:   List[str],
    ):
        """
        Iterates over classified paragraphs and adds each to the document
        using the correct formatting rules.
        """
        prev_type     = PARA_EMPTY
        chapter_count = 0
        is_first_body_after_chapter = False

        for item in paragraphs:
            ptype = item.ptype
            text  = item.cleaned.strip()

            if ptype == PARA_EMPTY:
                continue  # Whitespace management is handled by template spacing rules

            if ptype == PARA_FRONT_MATTER:
                continue  # Re-created on the title page

            elif ptype == PARA_CHAPTER:
                is_first_chapter = (chapter_count == 0)
                self._add_chapter_heading(doc, text, is_first_chapter)
                chapter_count += 1
                is_first_body_after_chapter = True

            elif ptype == PARA_SCENE_BREAK:
                self._add_scene_break(doc)
                is_first_body_after_chapter = False  # Scene break resets the opener rule

            else:  # PARA_BODY
                if not text:
                    continue
                opener = is_first_body_after_chapter
                self._add_body_paragraph(doc, text, is_chapter_opener=opener)
                is_first_body_after_chapter = False

            prev_type = ptype

        if chapter_count == 0:
            warnings.append(
                "No chapter headings were detected. The full text was formatted as body copy. "
                "If your document has chapters, ensure headings follow a recognisable pattern "
                "such as 'Chapter 1' or 'Chapter One' on their own line."
            )

    # ── Chapter Heading ───────────────────────────────────────────────────────

    def _add_chapter_heading(self, doc: Document, text: str, is_first: bool):
        t = self.template

        # All chapters except the first get a page break before them
        if not is_first:
            doc.add_page_break()

        # Push heading down if template calls for it
        if t.chapter_start_position == "third_down":
            for _ in range(t.chapter_blank_lines_above):
                blank = doc.add_paragraph()
                blank.paragraph_format.space_before = Pt(0)
                blank.paragraph_format.space_after  = Pt(0)
                blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Heading paragraph
        p = doc.add_paragraph()
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if t.chapter_centered else WD_ALIGN_PARAGRAPH.LEFT
        )
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(t.chapter_font_size_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        display = text.upper() if t.chapter_all_caps else text
        font_name = t.chapter_font_name or t.font_name

        run = p.add_run(display)
        run.font.name   = font_name
        run.font.size   = Pt(t.chapter_font_size_pt)
        run.bold        = t.chapter_bold
        run.italic      = t.chapter_italic

    # ── Scene Break ───────────────────────────────────────────────────────────

    def _add_scene_break(self, doc: Document):
        t = self.template
        gap = Pt(t.font_size_pt * 1.5)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = gap
        p.paragraph_format.space_after  = gap
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Inches(0)

        run = p.add_run(t.scene_break_text)
        run.font.name = t.font_name
        run.font.size = Pt(t.font_size_pt)

    # ── Body Paragraph ────────────────────────────────────────────────────────

    def _add_body_paragraph(self, doc: Document, text: str, is_chapter_opener: bool):
        """
        Standard body paragraph.
        is_chapter_opener: the first paragraph after a chapter heading has no
        first-line indent — this is the correct international typographic rule,
        used by every major publisher worldwide.
        """
        t = self.template
        p = doc.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY if t.body_justified
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        p.paragraph_format.space_before = Pt(t.space_before_para_pt)
        p.paragraph_format.space_after  = Pt(t.space_after_para_pt)
        self._set_line_spacing(p.paragraph_format, t.line_spacing)

        if is_chapter_opener:
            p.paragraph_format.first_line_indent = Inches(0)
        else:
            p.paragraph_format.first_line_indent = Inches(t.first_line_indent_in)

        run = p.add_run(text)
        run.font.name = t.font_name
        run.font.size = Pt(t.font_size_pt)


# ── Utility ───────────────────────────────────────────────────────────────────

def count_words_in_docx(path: str) -> int:
    """Returns the total word count of a .docx file."""
    doc   = Document(path)
    total = sum(len(p.text.split()) for p in doc.paragraphs)
    # Also count words inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += len(para.text.split())
    return total


def extract_raw_paragraphs(path: str) -> List[str]:
    """
    Reads every paragraph from a .docx file and returns raw text strings.
    Handles merged runs and table-embedded text.
    """
    doc = Document(path)
    paragraphs = []

    for para in doc.paragraphs:
        # Join all runs — avoids split-word artefacts from mixed formatting
        text = "".join(run.text for run in para.runs) or para.text
        paragraphs.append(text)

    # Some authors inadvertently use tables as layout tools
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)

    return paragraphs
