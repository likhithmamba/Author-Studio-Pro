import os
import sys
from pathlib import Path

# Add directories to path
sys.path.insert(0, str(Path(__file__).parent.parent / "raw claude novel editor pro" / "author_studio"))

from formatter import NovelFormatter, count_words_in_docx, extract_raw_paragraphs
from parser import ManuscriptParser
from templates import TEMPLATES

def test_run():
    # Make a dummy docx
    import docx
    doc = docx.Document()
    doc.add_paragraph("Chapter 1: The Beginning")
    doc.add_paragraph("This is a test paragraph.")
    test_docx_path = "test_dummy.docx"
    doc.save(test_docx_path)
    
    try:
        word_count = count_words_in_docx(test_docx_path)
        raw_paras = extract_raw_paragraphs(test_docx_path)
        
        parser = ManuscriptParser(api_key=None, model="mistral")
        parsed, learned = parser.parse(raw_paras)
        
        tpl = TEMPLATES["traditional"]
        fmt = NovelFormatter(tpl, author="Test Author", title="Test Title", word_count=word_count)
        
        output_path = "test_out.docx"
        fmt.build(parsed, output_path)
        print("SUCCESS! File formatted successfully.")
    except Exception as e:
        print(f"FAILED: {e}")
        import traceback
        traceback.print_exc()
        
if __name__ == "__main__":
    test_run()
